"""
PDF → Word (DOCX) Converter.
- Text tiếng Việt dùng get_text("dict") để PyMuPDF tự merge glyph.
- Bảng biểu dùng word-level spatial lookup để có khoảng trắng đúng.
- Dòng trong cùng block ghép bằng space (soft-wrap), không dùng \n cứng.
- OCR cho trang scan: Tesseract (vie+eng) → RapidOCR fallback.
"""

import io
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

GLOBAL_OCR_ENGINE = None
TESSERACT_AVAILABLE = None  # None = unchecked, True/False after first check


def _try_tesseract_ocr(pix_bytes: bytes, zoom: float) -> list:
    """
    Dùng Tesseract với vie+eng, group words theo (block_num, par_num, line_num),
    trả về list (x0,y0,x1,y1,text) theo tọa độ PDF gốc (đã chia zoom).
    """
    global TESSERACT_AVAILABLE
    if TESSERACT_AVAILABLE is False:
        return []
    try:
        import pytesseract
        from PIL import Image

        pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
        img = Image.open(io.BytesIO(pix_bytes))
        data = pytesseract.image_to_data(
            img, lang="vie+eng",
            output_type=pytesseract.Output.DICT,
            config="--psm 6",
        )
        TESSERACT_AVAILABLE = True

        # Group words by (block_num, par_num, line_num)
        lines: dict = {}
        n = len(data["text"])
        for i in range(n):
            txt = (data["text"][i] or "").strip()
            if not txt or int(data["conf"][i]) < 30:
                continue
            key = (data["block_num"][i], data["par_num"][i], data["line_num"][i])
            x, y, w, h = data["left"][i], data["top"][i], data["width"][i], data["height"][i]
            lines.setdefault(key, {"words": [], "x0": x, "y0": y, "x1": x + w, "y1": y + h})
            entry = lines[key]
            entry["words"].append(txt)
            entry["x0"] = min(entry["x0"], x)
            entry["y0"] = min(entry["y0"], y)
            entry["x1"] = max(entry["x1"], x + w)
            entry["y1"] = max(entry["y1"], y + h)

        results = []
        for key in sorted(lines.keys()):
            e = lines[key]
            text = " ".join(e["words"])
            results.append((
                e["x0"] / zoom, e["y0"] / zoom,
                e["x1"] / zoom, e["y1"] / zoom,
                text,
            ))
        return results
    except Exception:
        TESSERACT_AVAILABLE = False
        return []


# ── helpers ───────────────────────────────────────────────────────────────────

def _inside_table(bbox, table_bboxes):
    cx = (bbox[0] + bbox[2]) / 2
    cy = (bbox[1] + bbox[3]) / 2
    for tx0, ty0, tx1, ty1 in table_bboxes:
        if tx0 - 4 <= cx <= tx1 + 4 and ty0 - 4 <= cy <= ty1 + 4:
            return True
    return False


def _detect_align(x0, x1, page_width):
    mid = page_width / 2
    block_w = x1 - x0
    sym = abs(x0 - (page_width - x1))
    if x0 > mid * 1.1 and block_w < page_width * 0.45:
        return "right"
    if sym < page_width * 0.06 and block_w < page_width * 0.75:
        return "center"
    return "left"


def _words_in_rect(page_words, rx0, ry0, rx1, ry1):
    """Lấy words nằm trong rect, sắp theo y rồi x, ghép bằng space."""
    found = []
    for wx0, wy0, wx1, wy1, wtext, *_ in page_words:
        wcx = (wx0 + wx1) / 2
        wcy = (wy0 + wy1) / 2
        if rx0 - 2 <= wcx <= rx1 + 2 and ry0 - 2 <= wcy <= ry1 + 2:
            found.append((wy0, wx0, wtext))
    found.sort()
    return " ".join(w[2] for w in found)


def _merge_paragraph_blocks(blocks, page_width):
    """
    Gộp các block liền kề thực chất là cùng một đoạn văn.
    PDF thường lưu mỗi dòng hiển thị thành một block riêng.
    Điều kiện gộp: cùng align, gap dọc nhỏ, cùng indent, cỡ chữ tương đương.
    """
    if not blocks:
        return blocks

    merged = []
    current = blocks[0]

    for nxt in blocks[1:]:
        # Chỉ gộp 2 block đều là paragraph
        if current["type"] != "paragraph" or nxt["type"] != "paragraph":
            merged.append(current)
            current = nxt
            continue

        cb = current["bbox"]
        nb = nxt["bbox"]

        gap = nb[1] - cb[3]                        # khoảng cách dọc giữa 2 block
        line_h = max((cb[3] - cb[1]), 6)           # chiều cao dòng ước lượng
        cur_size = current.get("font_size", 13) or 13
        nxt_size = nxt.get("font_size", 13) or 13

        gap_ok       = -2 <= gap < line_h * 1.5    # âm một chút cho phép overlap nhẹ
        align_ok     = current.get("align") == nxt.get("align")
        indent_ok    = abs(cb[0] - nb[0]) < page_width * 0.07
        size_ok      = abs(cur_size - nxt_size) < 2.5
        # Không gộp nếu block tiếp theo bắt đầu bằng chữ hoa đứng đầu
        # sau dấu câu kết thúc (dấu chấm, chấm than, chấm hỏi) → đoạn mới
        cur_lines = current.get("lines", [])
        last_text = ""
        if cur_lines and cur_lines[-1].get("runs"):
            last_text = "".join(r["text"] for r in cur_lines[-1]["runs"]).rstrip()
        ends_sentence = last_text.endswith((".", "!", "?", ".\n"))

        should_merge = gap_ok and align_ok and indent_ok and size_ok and not ends_sentence

        if should_merge:
            nxt_lines = nxt.get("lines", [])
            # Đảm bảo có space giữa cuối block cũ và đầu block mới
            if cur_lines and cur_lines[-1].get("runs"):
                last_run = cur_lines[-1]["runs"][-1]
                if last_run["text"] and not last_run["text"].endswith(" "):
                    last_run["text"] += " "

            current = {
                **current,
                "bbox": [min(cb[0], nb[0]), cb[1], max(cb[2], nb[2]), nb[3]],
                "lines": cur_lines + nxt_lines,
                "font_size": max(cur_size, nxt_size),
                "is_bold": current.get("is_bold", False) or nxt.get("is_bold", False),
                "is_heading": current.get("is_heading", False) and nxt.get("is_heading", False),
            }
        else:
            merged.append(current)
            current = nxt

    merged.append(current)
    return merged


# ── core extraction ───────────────────────────────────────────────────────────

def parse_pdf_to_blocks(pdf_bytes: bytes) -> dict:
    global GLOBAL_OCR_ENGINE
    pdf_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    pages_data = []

    for page_index in range(len(pdf_doc)):
        page = pdf_doc.load_page(page_index)
        page_width = page.rect.width

        # 1. Table detection — dùng word-level lookup để tránh chữ dính
        page_words = page.get_text("words")  # (x0,y0,x1,y1,text,blk,ln,wn)
        tables = page.find_tables()
        table_bboxes = [list(t.bbox) for t in tables]
        extracted_tables = []

        for t in tables:
            row_count = t.row_count
            col_count = t.col_count
            # table.cells là list Rect|None, len = row_count * col_count
            cells_better = []
            for r in range(row_count):
                row = []
                for c in range(col_count):
                    cell_rect = t.cells[r * col_count + c]
                    if cell_rect is not None:
                        # Dùng word lookup để có khoảng trắng đúng
                        txt = _words_in_rect(page_words,
                                             cell_rect[0], cell_rect[1],
                                             cell_rect[2], cell_rect[3])
                    else:
                        # Merged cell hoặc None → fallback extract
                        raw = t.extract()
                        txt = str(raw[r][c] or "") if raw and r < len(raw) and c < len(raw[r]) else ""
                    row.append(txt)
                cells_better.append(row)

            extracted_tables.append({
                "type": "table",
                "bbox": list(t.bbox),
                "cells": cells_better,
            })

        # 2. Smart OCR trigger
        raw_text = page.get_text("text").strip()
        has_visuals = len(page.get_images()) > 0 or len(page.get_drawings()) > 0
        needs_ocr = len(raw_text) < 30 and has_visuals

        paragraphs = []

        if needs_ocr:
            zoom = 2.0
            pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
            pix_bytes = pix.tobytes("png")

            # Thử Tesseract trước (hỗ trợ tiếng Việt có dấu)
            tess_results = _try_tesseract_ocr(pix_bytes, zoom)

            if tess_results:
                for x0, y0, x1, y1, t_str in tess_results:
                    if _inside_table([x0, y0, x1, y1], table_bboxes):
                        continue
                    h = max(y1 - y0, 6)
                    size = round(h * 0.75, 1)
                    paragraphs.append({
                        "type": "paragraph",
                        "bbox": [x0, y0, x1, y1],
                        "align": _detect_align(x0, x1, page_width),
                        "is_bold": size > 14,
                        "is_heading": size > 14,
                        "font_size": size,
                        "lines": [{"runs": [{"text": t_str, "bold": size > 14,
                                             "italic": False, "size": size,
                                             "color": None}]}],
                    })
            else:
                # Fallback: RapidOCR
                global GLOBAL_OCR_ENGINE
                if GLOBAL_OCR_ENGINE is None:
                    from rapidocr_onnxruntime import RapidOCR
                    GLOBAL_OCR_ENGINE = RapidOCR()
                ocr_result, _ = GLOBAL_OCR_ENGINE(pix_bytes)
                if ocr_result:
                    for box, text, _ in ocr_result:
                        t_str = text.strip()
                        if not t_str:
                            continue
                        xs = [p[0] for p in box]
                        ys = [p[1] for p in box]
                        x0, y0 = min(xs) / zoom, min(ys) / zoom
                        x1, y1 = max(xs) / zoom, max(ys) / zoom
                        if _inside_table([x0, y0, x1, y1], table_bboxes):
                            continue
                        h = max(y1 - y0, 6)
                        size = round(h * 0.75, 1)
                        paragraphs.append({
                            "type": "paragraph",
                            "bbox": [x0, y0, x1, y1],
                            "align": _detect_align(x0, x1, page_width),
                            "is_bold": size > 14,
                            "is_heading": size > 14,
                            "font_size": size,
                            "lines": [{"runs": [{"text": t_str, "bold": size > 14,
                                                 "italic": False, "size": size,
                                                 "color": None}]}],
                        })
            paragraphs = _merge_paragraph_blocks(paragraphs, page_width)

        else:
            raw = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)

            # Tính heading threshold
            all_sizes = []
            for blk in raw["blocks"]:
                if blk["type"] == 0:
                    for ln in blk.get("lines", []):
                        for sp in ln["spans"]:
                            if sp["text"].strip():
                                all_sizes.append(sp["size"])
            avg_size = sum(all_sizes) / len(all_sizes) if all_sizes else 11.0
            heading_thresh = avg_size * 1.18

            for blk in raw["blocks"]:
                bbox = list(blk["bbox"])
                if _inside_table(bbox, table_bboxes):
                    continue

                if blk["type"] == 1:  # Image block
                    paragraphs.append({
                        "type": "image",
                        "bbox": bbox,
                        "ext": blk.get("ext", "png"),
                        "image": blk["image"],
                        "align": _detect_align(bbox[0], bbox[2], page_width),
                    })
                    continue

                if blk["type"] != 0:
                    continue

                # Text block
                line_list = []
                max_size = 0.0
                blk_bold = False

                for ln in blk["lines"]:
                    line_runs = []
                    for sp in ln["spans"]:
                        t = sp["text"]
                        if not t:
                            continue
                        sz = sp["size"]
                        bold = bool(sp["flags"] & 16)
                        italic = bool(sp["flags"] & 2)
                        # Extract color — PyMuPDF stores as 0xRRGGBB int
                        raw_color = sp.get("color", 0)
                        color_hex = None if raw_color == 0 else f"#{raw_color:06X}"
                        if sz > max_size:
                            max_size = sz
                        if bold:
                            blk_bold = True
                        line_runs.append({
                            "text": t,
                            "bold": bold,
                            "italic": italic,
                            "size": round(sz, 1),
                            "color": color_hex,
                        })

                    if line_runs:
                        line_list.append({"runs": line_runs})

                if not line_list:
                    continue

                x0, y0, x1, y1 = bbox
                is_heading = (max_size >= heading_thresh) or blk_bold

                paragraphs.append({
                    "type": "paragraph",
                    "bbox": bbox,
                    "align": _detect_align(x0, x1, page_width),
                    "is_bold": blk_bold,
                    "is_heading": is_heading,
                    "font_size": round(max_size, 1),
                    "lines": line_list,
                })

        # Gộp các block text liền kề thuộc cùng đoạn văn
        paragraphs = _merge_paragraph_blocks(paragraphs, page_width)

        all_blocks = paragraphs + extracted_tables
        all_blocks.sort(key=lambda b: b["bbox"][1])
        pages_data.append({"page_num": page_index + 1, "blocks": all_blocks})

    pdf_doc.close()
    return {"pages": pages_data}


# ── DOCX builder ──────────────────────────────────────────────────────────────

def build_docx_from_blocks(blocks_data: dict) -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.18)
        section.right_margin  = Inches(1.18)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(13)
    normal.paragraph_format.space_after  = Pt(4)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.line_spacing = 1.15

    def _add_para(block):
        align_str  = block.get("align", "left")
        dom_size   = block.get("font_size", 13.0) or 13.0
        lines      = block.get("lines", [])
        is_heading = block.get("is_heading", False)
        blk_bold   = block.get("is_bold", False)
        bbox       = block.get("bbox", [0, 0, 0, 0])

        p = doc.add_paragraph()
        if align_str == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align_str == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            x0 = bbox[0]
            if x0 > 95:
                p.paragraph_format.left_indent = Inches(
                    min(2.5, max(0.0, (x0 - 72) / 72.0))
                )

        # Gộp các dòng trong cùng block bằng SPACE, không dùng \n cứng.
        # PDF soft-wrap theo cột → Word tự reflow theo lề trang.
        all_runs = []
        for l_idx, ln in enumerate(lines):
            for r_idx, rdata in enumerate(ln.get("runs", [])):
                txt = rdata.get("text", "")
                if not txt:
                    continue

                # Đảm bảo có space giữa các dòng (không phải giữa các span cùng dòng)
                if r_idx == 0 and l_idx > 0:
                    # Đây là run đầu của dòng mới — thêm space nối với dòng trước
                    if all_runs and not all_runs[-1]["text"].endswith(" "):
                        # Thêm space vào run trước hoặc tạo run space mới
                        all_runs[-1]["text"] += " "

                all_runs.append({
                    "text": txt,
                    "bold": rdata.get("bold", False) or (is_heading and blk_bold),
                    "italic": rdata.get("italic", False),
                    "size": rdata.get("size") or dom_size,
                })

        for rdata in all_runs:
            txt = rdata["text"]
            if not txt:
                continue
            run = p.add_run(txt)
            run.font.name = "Times New Roman"
            run.font.size = Pt(round(rdata["size"], 1))
            run.bold   = rdata["bold"]
            run.italic = rdata["italic"]
            color_hex = rdata.get("color")
            if color_hex:
                r = int(color_hex[1:3], 16)
                g = int(color_hex[3:5], 16)
                b = int(color_hex[5:7], 16)
                run.font.color.rgb = RGBColor(r, g, b)

    for p_idx, page in enumerate(blocks_data.get("pages", [])):
        for block in page.get("blocks", []):
            if block["type"] == "table":
                cells = block.get("cells", [])
                if not cells:
                    continue
                n_rows = len(cells)
                n_cols = max((len(r) for r in cells), default=0)
                if n_cols == 0:
                    continue

                tbl = doc.add_table(rows=n_rows, cols=n_cols)
                tbl.style = "Table Grid"
                for r_i, row in enumerate(cells):
                    for c_i in range(n_cols):
                        val = str(row[c_i] or "") if c_i < len(row) else ""
                        cell = tbl.cell(r_i, c_i)
                        cell.text = val
                        for cp in cell.paragraphs:
                            for cr in cp.runs:
                                cr.font.name = "Times New Roman"
                                cr.font.size = Pt(11)

            elif block["type"] == "image":
                img_bytes = block.get("image")
                if not img_bytes:
                    continue
                p = doc.add_paragraph()
                align_str = block.get("align", "center")
                p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if align_str == "center"
                               else WD_ALIGN_PARAGRAPH.RIGHT if align_str == "right"
                               else WD_ALIGN_PARAGRAPH.LEFT)
                pdf_width = block["bbox"][2] - block["bbox"][0]
                docx_width = Inches(min(5.8, max(0.6, pdf_width / 72.0)))
                run = p.add_run()
                try:
                    run.add_picture(io.BytesIO(img_bytes), width=docx_width)
                except Exception:
                    run.text = f"[Hình ảnh: lỗi định dạng .{block.get('ext', 'img')}]"

            else:
                _add_para(block)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


# ── convenience ───────────────────────────────────────────────────────────────

def convert_pdf_bytes_to_docx_bytes(pdf_bytes: bytes) -> bytes:
    return build_docx_from_blocks(parse_pdf_to_blocks(pdf_bytes))
